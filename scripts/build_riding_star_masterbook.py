from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "dist"
DOCX_PATH = OUT_DIR / "Riding-star_방송시나리오_마스터북.docx"
MD_PATH = OUT_DIR / "Riding-star_방송시나리오_마스터북.md"

ACCENT = "1F6F73"
ACCENT_DARK = "174A52"
ACCENT_LIGHT = "EAF6F5"
SOFT_GRAY = "F4F6F8"
BORDER = "D2DADF"
TEXT = "1E2930"


def set_east_asia_font(run, font_name: str = "Malgun Gothic") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name: str = "Malgun Gothic", size_pt: float | None = None) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_cell(cell, fill: str | None = None, bold: bool = False, color: str = TEXT, size: float = 9.0,
               align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.12
        for run in paragraph.runs:
            set_east_asia_font(run)
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.0, color: str = TEXT,
                  align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    lines = text.split("\n")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_east_asia_font(run)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def add_small_label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)


def add_section_title(document: Document, title: str, subtitle: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    set_east_asia_font(run)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    if subtitle:
        sub = document.add_paragraph()
        sub.paragraph_format.space_after = Pt(8)
        sub_run = sub.add_run(subtitle)
        set_east_asia_font(sub_run)
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor.from_string("5D6B72")


def add_note_box(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, ACCENT_LIGHT)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    set_cell_borders(cell, color="B5DAD7")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    title_run = paragraph.add_run(title)
    set_east_asia_font(title_run)
    title_run.font.size = Pt(9.5)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    body_p.paragraph_format.line_spacing = 1.18
    body_run = body_p.add_run(body)
    set_east_asia_font(body_run)
    body_run.font.size = Pt(9)
    body_run.font.color.rgb = RGBColor.from_string(TEXT)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None,
              font_size: float = 8.7, header_size: float = 8.7, header_fill: str = ACCENT) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False if widths_cm else True
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr_cells[i], header, bold=True, size=header_size, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        style_cell(hdr_cells[i], fill=header_fill, bold=True, color="FFFFFF", size=header_size,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths_cm:
            hdr_cells[i].width = Cm(widths_cm[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size,
                          align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 2, len(row) - 1) else WD_ALIGN_PARAGRAPH.LEFT)
            style_cell(cells[i], fill=None, size=font_size,
                       align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 2, len(row) - 1) else WD_ALIGN_PARAGRAPH.LEFT)
            if widths_cm:
                cells[i].width = Cm(widths_cm[i])
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_two_col_meta(document: Document, rows: list[tuple[str, str, str, str]]) -> None:
    table = document.add_table(rows=len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    widths = [Cm(2.7), Cm(6.0), Cm(2.7), Cm(6.0)]
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, bold=c_idx in (0, 2), size=8.8,
                          color="FFFFFF" if c_idx in (0, 2) else TEXT,
                          align=WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT)
            style_cell(cell, fill=ACCENT if c_idx in (0, 2) else None, bold=c_idx in (0, 2),
                       color="FFFFFF" if c_idx in (0, 2) else TEXT, size=8.8,
                       align=WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT)
            cell.width = widths[c_idx]
    document.add_paragraph().paragraph_format.space_after = Pt(3)


def add_script_block(document: Document, title: str, purpose: str, rows: list[tuple[str, str, str]]) -> None:
    add_section_title(document, title, purpose)
    for speaker, line, note in rows:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.25)
        paragraph.paragraph_format.first_line_indent = Cm(-0.25)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.22

        speaker_run = paragraph.add_run(f"{speaker}: ")
        set_east_asia_font(speaker_run)
        speaker_run.font.size = Pt(10.2)
        speaker_run.font.bold = True
        speaker_run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

        line_run = paragraph.add_run(line)
        set_east_asia_font(line_run)
        line_run.font.size = Pt(10.2)
        line_run.font.color.rgb = RGBColor.from_string(TEXT)

        if note:
            note_run = paragraph.add_run(f"  [{note}]")
            set_east_asia_font(note_run)
            note_run.font.size = Pt(8.2)
            note_run.font.color.rgb = RGBColor.from_string("6A7880")

    document.add_paragraph().paragraph_format.space_after = Pt(3)


def add_markdown_heading(lines: list[str], level: int, title: str) -> None:
    lines.append(f"{'#' * level} {title}")
    lines.append("")


def build_markdown() -> str:
    lines: list[str] = []
    add_markdown_heading(lines, 1, "Riding-star 방송 시나리오 마스터북")
    lines.append("자전거 이야기를 천천히, 같이, 멀리 전하는 라디오 방송 구성안")
    lines.append("")
    add_markdown_heading(lines, 2, "회차별 모음")
    lines.append("| 회차 | 주제 | 게스트 | 녹음일 | 방송일 | 키워드 | 추천 음악 | 자료 위치 |")
    lines.append("|---|---|---|---|---|---|---|---|")
    for idx in range(1, 9):
        if idx == 1:
            lines.append("| EP.01 | 천천히, 같이, 멀리 | 이은영 | 2025.08.28 | 2025.09.__ | 첫 라이딩, 동반자, 장거리 | 1~4곡 | 본 문서 |")
        else:
            lines.append(f"| EP.{idx:02d} |  |  |  |  |  |  |  |")
    lines.append("")
    add_markdown_heading(lines, 2, "전체 방송 시간표")
    lines.append("| 구성 | 시간 | 누적 | 세부사항 | 출연 |")
    lines.append("|---|---:|---:|---|---|")
    for row in RUNNING_ROWS:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3].replace(chr(10), '<br>')} | {row[4]} |")
    lines.append("")
    for title, purpose, rows in SCRIPT_BLOCKS:
        add_markdown_heading(lines, 2, title)
        lines.append(purpose)
        lines.append("")
        for speaker, line, note in rows:
            prefix = f"**{speaker}**"
            note_text = f" ({note})" if note else ""
            lines.append(f"- {prefix}: {line}{note_text}")
        lines.append("")
    add_markdown_heading(lines, 2, "게스트 추천 음악 기록")
    lines.append("| 순서 | 연결 구간 | 곡명 | 아티스트 | 추천 이유 | 큐시트 메모 |")
    lines.append("|---|---|---|---|---|---|")
    for row in MUSIC_ROWS:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} | {row[5]} |")
    lines.append("")
    return "\n".join(lines)


RUNNING_ROWS = [
    ["오프닝", "3:00", "00:00-03:00",
     "시그널 음악\n진행자 인사 및 코너 설명\n오늘의 키워드: 천천히, 같이, 멀리\n게스트 소개로 자연스럽게 연결", "진행자"],
    ["음악 1", "4:00", "03:00-07:00",
     "게스트 추천곡 1\n첫 라이딩 기억으로 넘어가는 분위기 만들기", "진행자"],
    ["토크 1", "10:00", "07:00-17:00",
     "처음 자전거를 탔던 날\n페달, 바람, 넘어짐, 첫 성공의 감각\n어릴 때와 지금의 라이딩 감각 비교", "진행자 & 게스트"],
    ["음악 2", "3:00", "17:00-20:00",
     "게스트 추천곡 2\n천천히 달리는 장면을 떠올리게 하는 곡", "진행자"],
    ["토크 2", "15:00", "20:00-35:00",
     "천천히 같이 달리는 이유\n동반 라이딩에서 생기는 대화와 배려\n속도보다 중요한 리듬, 안전, 관계", "진행자 & 게스트"],
    ["음악 3", "3:00", "35:00-38:00",
     "게스트 추천곡 3\n장거리 라이딩 전환용 브리지", "진행자"],
    ["토크 3", "12:00", "38:00-50:00",
     "멀리 떠났던 라이딩 기억\n추천 코스와 코스 근처 맛집\n언젠가 꼭 가보고 싶은 장거리 목적지", "진행자 & 게스트"],
    ["마무리 질문", "3:00", "50:00-53:00",
     "오늘 방송 후 어떤 하루를 보내고 싶은지\n게스트가 진행자에게 묻는 역질문\n오늘의 한 문장 정리", "진행자 & 게스트"],
    ["엔딩 멘트", "3:00", "53:00-56:00",
     "게스트 감사 인사\n청취자에게 다음 회차 예고\n시그널 음악 및 클로징", "진행자"],
]


SCRIPT_BLOCKS = [
    (
        "오프닝 멘트 (00:00-03:00)",
        "목표: 방송의 결을 밝히고, 오늘의 주제와 게스트를 부드럽게 소개한다.",
        [
            ("왕규", "안녕하세요. 바람은 살짝 선선하고 햇살은 한층 부드러워졌습니다. 오늘도 자전거 타기 좋은 시간, Riding-star 문을 열겠습니다.", "밝게"),
            ("정규", "오늘은 자전거를 빠르게 달리는 이야기보다, 천천히 같이 달리며 멀리 가는 이야기를 준비했습니다. 페달을 밟는 속도보다 마음의 리듬을 먼저 맞춰보려고 합니다.", ""),
            ("왕규", "라이딩을 하다 보면 풍경도 보고, 옆 사람의 숨소리도 듣고, 가끔은 맛있는 곳까지 찾아가게 되잖아요. 오늘은 그런 장면들을 하나씩 꺼내보겠습니다.", ""),
            ("정규", "오늘의 게스트를 소개합니다. 초등학교 3학년 때 아빠 자전거 안장에 앉아 처음 바람을 만났고, 지금은 즐겁고 건강한 라이딩을 이어가고 있는 이은영 님입니다.", "게스트 소개"),
            ("게스트", "안녕하세요. 이은영입니다. 불러주셔서 감사합니다.", "인사"),
            ("왕규", "은영 님과 함께 오늘의 세 단어, 천천히, 같이, 멀리를 따라가 보겠습니다. 먼저 게스트 추천곡 한 곡 듣고 첫 이야기로 돌아오겠습니다.", "음악 1 연결"),
        ],
    ),
    (
        "토크 1: 처음 자전거를 탔던 날 (07:00-17:00)",
        "목표: 게스트의 첫 자전거 기억을 통해 청취자가 자기 경험을 떠올리게 한다.",
        [
            ("왕규", "자전거를 처음 배울 때는 왜 그렇게 온몸에 힘이 들어갔을까요. 은영 님은 처음 페달을 밟던 날을 기억하시나요?", "질문 1"),
            ("게스트", "어릴 때 아빠가 잡아주시던 기억이 나요. 중심을 잡는 게 무섭기도 했지만, 어느 순간 혼자 앞으로 나가고 있어서 신기했어요.", "답변 예시"),
            ("정규", "그때 제일 크게 남은 감각은 뭐였나요. 바람, 무서움, 웃음, 아니면 넘어질까 봐 긴장하던 마음 중에요.", "꼬리 질문"),
            ("게스트", "바람이 제일 기억나요. 빠르지 않았는데도 제가 직접 앞으로 간다는 느낌이 좋았어요.", ""),
            ("왕규", "처음 배울 때 넘어졌던 기억도 있을 것 같아요. 지금 생각하면 웃을 수 있는 장면이 있나요?", "장면 유도"),
            ("게스트", "넘어지고 나서도 다시 타고 싶었던 걸 보면, 그때 이미 자전거가 꽤 좋았던 것 같아요.", ""),
            ("정규", "그 마음이 지금 라이딩까지 이어졌네요. 첫 번째 이야기는 여기서 잠깐 쉬고, 천천히 달리는 장면에 어울리는 두 번째 음악 듣겠습니다.", "음악 2 연결"),
        ],
    ),
    (
        "토크 2: 천천히 같이 달리며 나누는 이야기 (20:00-35:00)",
        "목표: 함께 타는 라이딩의 배려, 대화, 안전 감각을 중심으로 이야기를 확장한다.",
        [
            ("정규", "오늘의 첫 번째 키워드는 천천히입니다. 자전거를 타다 보면 일부러 속도를 낮추는 순간이 있잖아요. 은영 님은 언제 천천히 달리게 되나요?", "질문 1"),
            ("게스트", "풍경이 좋거나 같이 타는 사람이 있을 때요. 천천히 가야 보이는 게 많더라고요.", ""),
            ("왕규", "함께 달릴 때는 속도를 맞추는 것도 기술 같아요. 남편분이나 지인과 같이 탈 때 서로 맞춰가는 방식이 있나요?", "관계 질문"),
            ("게스트", "앞에서 너무 멀리 가지 않기, 뒤를 자주 확인하기, 힘든 사람을 기다려주기 같은 것들이요.", ""),
            ("정규", "라이딩 모임에서 뒤를 지켜주는 역할을 즐긴다고 하셨는데, 뒤에서 보면 앞에 있을 때와 다른 풍경이 보이나요?", "배려 질문"),
            ("게스트", "네. 뒤에 있으면 누가 힘들어하는지, 어디서 속도를 줄여야 하는지가 더 잘 보여요.", ""),
            ("왕규", "자전거는 혼자 타도 좋지만 같이 타면 이야기가 생깁니다. 은영 님의 플레이리스트에서 같이 달릴 때 어울리는 곡 하나 더 들어볼까요?", "음악 3 연결"),
        ],
    ),
    (
        "토크 3: 멀리, 코스와 맛집 이야기 (38:00-50:00)",
        "목표: 장거리 라이딩의 설렘과 실제 정보를 함께 담아 청취자의 다음 라이딩을 자극한다.",
        [
            ("정규", "세 번째 키워드는 멀리입니다. 은영 님에게 가장 기억에 남는 장거리 라이딩은 어디였나요?", "질문 1"),
            ("게스트", "섬진강을 따라 갔던 1박 2일 여행이 기억에 남아요. 길도 좋고, 중간중간 쉬어가는 시간이 좋았어요.", "답변 예시"),
            ("왕규", "그 코스를 누군가에게 추천한다면 어느 구간을 꼭 달려보라고 말하고 싶으세요?", "정보 질문"),
            ("게스트", "강을 옆에 두고 달리는 구간이요. 속도를 내지 않아도 풍경이 계속 바뀌어서 지루하지 않아요.", ""),
            ("정규", "라이딩 이야기에서 맛집을 빼면 서운합니다. 그 코스 근처에서 기억나는 음식이나 쉬어가기 좋은 곳이 있었나요?", "맛집 질문"),
            ("게스트", "따뜻한 국물이나 간단히 먹기 좋은 분식집처럼 부담 없는 곳이 좋았어요. 오래 앉아 쉬는 것도 중요하더라고요.", ""),
            ("왕규", "언젠가 꼭 가보고 싶은 멀리의 목적지도 궁금합니다. 아직 실현하지 않았지만 마음속에 담아둔 코스가 있나요?", "꿈 질문"),
            ("게스트", "동해안 자전거길을 길게 달려보고 싶어요. 바다를 보면서 천천히 가보고 싶습니다.", ""),
        ],
    ),
    (
        "마무리 질문 & 게스트의 역질문 (50:00-53:00)",
        "목표: 방송의 핵심 감정을 한 번 더 정리하고, 게스트가 진행자에게 질문하는 장면을 만든다.",
        [
            ("정규", "마지막 질문입니다. 오늘 방송을 마치고 나면 어떤 하루로 마무리하고 싶으세요?", "마무리 질문"),
            ("게스트", "천천히 걸어가거나, 좋아하는 음악을 들으면서 오늘 나눈 이야기를 다시 떠올리고 싶어요.", ""),
            ("왕규", "이번에는 은영 님이 저희에게 질문을 하나 던져주셔도 좋습니다. 오늘의 주제, 천천히 같이 멀리에서 떠오른 질문이 있을까요?", "역질문 유도"),
            ("게스트", "두 분에게 자전거는 어떤 모습인가요? 오프닝만 들었을 때도 아직 잘 모르겠더라고요.", "역질문"),
            ("정규", "저에게 자전거는 조급한 마음을 내려놓게 하는 도구입니다. 천천히 가도 도착한다는 걸 몸으로 알려주는 쪽에 가까워요.", "진행자 답변"),
            ("왕규", "저에게는 같이 웃을 수 있는 핑계입니다. 커피 한 잔 마시고, 라디오 소리 듣고, 옆 사람과 같은 바람을 맞는 시간이라고 생각합니다.", "진행자 답변"),
        ],
    ),
    (
        "엔딩 멘트 (53:00-56:00)",
        "목표: 감사 인사, 주제 회수, 다음 회차 예고를 짧고 선명하게 마무리한다.",
        [
            ("왕규", "오늘 Riding-star는 천천히, 같이, 멀리라는 세 단어를 따라 달렸습니다. 숨은 이야기를 들려주신 이은영 님 고맙습니다.", "감사"),
            ("게스트", "저도 감사합니다. 듣는 분들도 마음속에서 페달을 천천히 밟아보시면 좋겠습니다.", ""),
            ("정규", "청취자 여러분, 이번 주에는 목적지를 조금 늦게 정하고 주변 풍경을 더 오래 바라보셔도 좋겠습니다.", "청취자 인사"),
            ("왕규", "다음 회차에는 또 다른 두 바퀴 이야기를 가지고 돌아오겠습니다. 지금까지 Riding-star였습니다.", "클로징"),
        ],
    ),
]


MUSIC_ROWS = [
    ["1", "오프닝 후", "게스트 추천곡 1", "아티스트명", "첫 라이딩의 설렘을 여는 곡", "03:00 인 / 07:00 아웃"],
    ["2", "토크 1 후", "게스트 추천곡 2", "아티스트명", "천천히 달리는 장면과 어울리는 곡", "17:00 인 / 20:00 아웃"],
    ["3", "토크 2 후", "게스트 추천곡 3", "아티스트명", "장거리 라이딩으로 넘어가는 브리지", "35:00 인 / 38:00 아웃"],
    ["4", "엔딩", "엔딩 추천곡", "아티스트명", "방송 여운을 남기는 곡", "56:00 이후"],
    ["예비", "필요 시", "", "", "방송 시간 변동 대비", ""],
]


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    styles = document.styles
    set_style_font(styles["Normal"], size_pt=9.6)
    styles["Normal"].paragraph_format.line_spacing = 1.18
    styles["Normal"].paragraph_format.space_after = Pt(5)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        set_style_font(styles[style_name])
        styles[style_name].font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    header = section.header.paragraphs[0]
    header.text = "Riding-star | 방송 시나리오 마스터북"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_east_asia_font(run)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("7A8890")

    footer = section.footer.paragraphs[0]
    footer.text = "전주공동체라디오 93.5MHz"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_east_asia_font(run)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("7A8890")


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(70)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Riding-star")
    set_east_asia_font(run)
    run.font.size = Pt(31)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    sub_run = subtitle.add_run("자전거 이야기를 천천히, 같이, 멀리 전하는 방송 시나리오 북")
    set_east_asia_font(sub_run)
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor.from_string("52616A")

    add_note_box(
        document,
        "문서 사용법",
        "첫 페이지에는 회차별 모음을 기록하고, 각 회차는 전체 시간표 -> 오프닝 -> 토크 1, 2, 3 -> 마무리 질문 -> 엔딩 -> 게스트 추천 음악 기록 순서로 복제해 사용합니다.",
    )

    add_section_title(document, "회차별 모음", "방송이 쌓일수록 한눈에 찾아보기 위한 인덱스입니다.")
    archive_rows = []
    for idx in range(1, 9):
        if idx == 1:
            archive_rows.append([f"EP.{idx:02d}", "천천히, 같이, 멀리", "이은영", "2025.08.28", "2025.09.__", "첫 라이딩 / 동반 라이딩 / 장거리", "추천곡 1~4", "본 문서"])
        else:
            archive_rows.append([f"EP.{idx:02d}", "", "", "", "", "", "", ""])
    add_table(
        document,
        headers=["회차", "주제", "게스트", "녹음일", "방송일", "핵심 키워드", "추천 음악", "자료 위치"],
        rows=archive_rows,
        widths_cm=[1.35, 3.1, 1.9, 2.25, 2.25, 3.2, 2.1, 1.9],
        font_size=7.7,
        header_size=7.6,
    )

    add_section_title(document, "제작 체크리스트", "녹음 전후로 빠르게 확인할 항목입니다.")
    checklist_rows = [
        ["녹음 전", "게스트 이름/호칭 확인, 오늘의 키워드 3개 확정, 추천곡 3~4곡 수집, 곡명/아티스트 표기 확인"],
        ["녹음 중", "각 파트 시작 시간 체크, 음악 인/아웃 큐 확인, 게스트 답변이 길어질 때 마무리 질문으로 정리"],
        ["편집 후", "실제 러닝타임 기록, 삽입곡 최종 확인, 회차별 모음 페이지 업데이트, 다음 회차 예고 문장 저장"],
    ]
    add_table(document, headers=["단계", "확인 내용"], rows=checklist_rows, widths_cm=[2.2, 15.2], font_size=8.5)

    document.add_page_break()

    add_section_title(document, "EP.01 전체 방송 시간표", "목표 러닝타임 56분. 실제 녹음 후 누적 시간만 수정하면 됩니다.")
    add_two_col_meta(
        document,
        [
            ("프로그램명", "Riding-star (천천히. 같이. 멀리.)", "주파수", "전주공동체라디오 93.5MHz"),
            ("진행", "왕규, 정규", "게스트", "이은영"),
            ("녹음장소", "방송실", "녹음일시", "2025.08.28"),
            ("방송일시", "2025.09.__  __:__", "회차 주제", "천천히, 같이, 멀리"),
        ],
    )
    add_table(
        document,
        headers=["구성", "소요", "누적", "세부사항", "출연"],
        rows=RUNNING_ROWS,
        widths_cm=[1.75, 1.35, 2.15, 9.7, 2.3],
        font_size=7.8,
        header_size=7.8,
    )

    document.add_page_break()

    for idx, (title_text, purpose, rows) in enumerate(SCRIPT_BLOCKS):
        if idx in (1, 3, 5):
            document.add_page_break()
        add_script_block(document, title_text, purpose, rows)

    document.add_page_break()

    add_section_title(document, "게스트 추천 음악 기록", "게스트가 추천한 음악, 연결 구간, 추천 이유를 남겨 다음 회차 자료로도 활용합니다.")
    add_table(
        document,
        headers=["순서", "연결 구간", "곡명", "아티스트", "추천 이유", "큐시트 메모"],
        rows=MUSIC_ROWS,
        widths_cm=[1.2, 2.25, 3.2, 2.7, 5.5, 2.65],
        font_size=8.2,
        header_size=8.0,
    )

    add_section_title(document, "방송 후 정리", "실제 방송이 끝난 뒤 회차별 모음에 옮겨 적을 핵심 기록입니다.")
    post_rows = [
        ["실제 러닝타임", "____분 ____초", "편집 메모", ""],
        ["가장 좋았던 장면", "", "삭제/축약한 부분", ""],
        ["청취자에게 남길 한 문장", "", "다음 회차 예고", ""],
        ["다음 게스트 후보", "", "다음 주제 후보", ""],
    ]
    add_table(
        document,
        headers=["항목", "내용", "항목", "내용"],
        rows=post_rows,
        widths_cm=[2.6, 6.1, 2.6, 6.1],
        font_size=8.5,
        header_size=8.2,
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    add_section_title(document, "다음 회차 복제용 빈 구성", "아래 항목만 채우면 새 회차 기본 구성이 완성됩니다.")
    add_two_col_meta(
        document,
        [
            ("회차", "EP.__", "주제", ""),
            ("진행", "", "게스트", ""),
            ("녹음일시", "", "방송일시", ""),
            ("오늘의 키워드", "1)  2)  3)", "목표 러닝타임", "56:00"),
        ],
    )
    template_rows = [
        ["오프닝", "00:00-03:00", "오늘의 주제 소개 / 게스트 소개", ""],
        ["음악 1", "03:00-07:00", "게스트 추천곡", ""],
        ["토크 1", "07:00-17:00", "개인 경험을 여는 질문", ""],
        ["음악 2", "17:00-20:00", "토크 2 전환곡", ""],
        ["토크 2", "20:00-35:00", "관계/가치/생활 이야기", ""],
        ["음악 3", "35:00-38:00", "토크 3 전환곡", ""],
        ["토크 3", "38:00-50:00", "코스/장소/정보 이야기", ""],
        ["마무리 질문", "50:00-53:00", "오늘의 한 문장 / 게스트 역질문", ""],
        ["엔딩", "53:00-56:00", "감사 인사 / 다음 회차 예고", ""],
    ]
    add_table(
        document,
        headers=["구성", "시간", "준비 내용", "수정 메모"],
        rows=template_rows,
        widths_cm=[2.1, 2.5, 8.7, 4.2],
        font_size=8.3,
        header_size=8.1,
    )

    document.save(DOCX_PATH)
    MD_PATH.write_text(build_markdown(), encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
